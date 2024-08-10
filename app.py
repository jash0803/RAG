from PyPDF2 import PdfReader
import streamlit as st
from streamlit_option_menu import option_menu
import pandas as pd
#import openpyxl
from docx.api import Document


st.set_page_config(page_title="Extract data from different file types",
                       page_icon=":books:")

with st.sidebar:
    selected = option_menu(
        'FILE EXTRACTOR',
        ['PDF','EXCEL','WORD'],
        default_index=0
    )


def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text


if selected == 'PDF':
    st.header("PDF EXTRACTOR")
    
    pdf_docs = st.file_uploader(
        "Upload your PDFs here and click on 'Process'", accept_multiple_files=True)
    if st.button("Process"):
        with st.spinner("Processing"):
            # get pdf text
            raw_text = get_pdf_text(pdf_docs)
            st.write(raw_text)

if selected == 'EXCEL':
    st.header("EXCEL EXTRACTOR")
    
    # Allow multiple file uploads
    excel_docs = st.file_uploader("Upload your EXCELs here and click on 'Process'", accept_multiple_files=True)

    if st.button("Process"):
        if excel_docs is not None:
            for excel_file in excel_docs:
                # Read each Excel file
                #openpyxl is used to read 2010 excel format which the most common one
                df = pd.read_excel(excel_file, engine='openpyxl')
                st.write(f"Processed file: {excel_file.name}")
                st.write(df)
        else:
            st.warning("Please upload at least one Excel file.")


if selected == 'WORD':
    st.header("WORD EXTRACTOR")
    
    # Allow multiple file uploads
    word_docs = st.file_uploader("Upload your WORDs here and click on 'Process'", accept_multiple_files=True)

    if st.button("Process"):
        if word_docs is not None:
            for word_file in word_docs:
                # Read each WORD file
                word_doc = Document(word_file)
                textdata=""
                for p in word_doc.paragraphs:
                    textdata+=p.text+"\n"
                st.write(textdata)

                # Initialize an empty string to store table data
                tabledata = ""

                # Loop through each table in the document
                for table in word_doc.tables:
                    for row in table.rows:
                        # Loop through each cell in the row and concatenate the text with "|"
                        for cell in row.cells:
                            tabledata += cell.text + "|"
                        # Add a newline at the end of each row
                        tabledata += "\n"

                # Display the extracted table data using Streamlit
                st.write(tabledata)

        else:
            st.warning("Please upload at least one WORD file.")